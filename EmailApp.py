import customtkinter as ctk
from tkinter import filedialog
import pandas as pd
import smtplib
import re
import os
import threading
import logging
from email.utils import parseaddr
from typing import Optional, List, Any, Tuple
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dotenv import load_dotenv
from docx import Document
import openpyxl
from datetime import datetime
import webbrowser
import tempfile
from pathlib import Path
import html as _html

load_dotenv()


SMTP_SERVER: Optional[str] = os.getenv("SMTP_SERVER")
try:
    SMTP_PORT: int = int(os.getenv("SMTP_PORT", "587"))
except ValueError:
    SMTP_PORT = 587
SMTP_USER: Optional[str] = os.getenv("SMTP_USER")
SMTP_PASSWORD: Optional[str] = os.getenv("SMTP_PASSWORD")
USE_TLS: bool = os.getenv("USE_TLS", "True").lower() in ("1", "true", "yes")

# basic logging to file (keeps existing email_log.txt behavior consistent)
logging.basicConfig(
    filename="email_log.txt",
    level=logging.INFO,
    format="[%(asctime)s] %(message)s",
    datefmt="%d-%m-%Y %H:%M:%S",
)

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("dark-blue")


class EmailApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.iconbitmap("mail.ico")

        self.title("Mail Merge - Excel Sender")
        self.geometry("1000x700")

        self.df: Optional[pd.DataFrame] = None
        self.file_path: Optional[str] = None

        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(3, weight=1)
        self._last_preview_path: Optional[Path] = None

        # Sidebar
        self.sidebar = ctk.CTkFrame(self, width=200)
        self.sidebar.grid(row=0, column=0, rowspan=4, sticky="nsw", padx=10, pady=10)

        self.load_btn = ctk.CTkButton(
            self.sidebar, text="Carica Excel", command=self.load_file
        )
        self.load_btn.pack(pady=10, fill="x")

        self.load_word_btn = ctk.CTkButton(
            self.sidebar, text="Carica Template Word", command=self.load_word_template
        )
        self.load_word_btn.pack(pady=10, fill="x")

        self.create_word_btn = ctk.CTkButton(
            self.sidebar, text="Crea Template Word", command=self.create_word_template
        )
        self.create_word_btn.pack(pady=10, fill="x")

        self.create_excel_btn = ctk.CTkButton(
            self.sidebar, text="Crea Template Excel", command=self.create_excel_template
        )
        self.create_excel_btn.pack(pady=10, fill="x")

        self.preview_btn = ctk.CTkButton(
            self.sidebar, text="Anteprima Prima Email", command=self.preview_email
        )
        self.preview_btn.pack(pady=10, fill="x")

        self.fields_btn = ctk.CTkButton(
            self.sidebar, text="Mostra Campi", command=self.show_fields
        )
        self.fields_btn.pack(pady=10, fill="x")

        self.test_mode = ctk.CTkCheckBox(
            self.sidebar, text="Modalità TEST (invio solo a me)"
        )
        self.test_mode.pack(pady=20)

        self.send_btn = ctk.CTkButton(
            self.sidebar, text="Invia Email", command=self.send_emails
        )
        self.send_btn.pack(pady=20, fill="x")

        self.info_btn = ctk.CTkButton(
            self.sidebar, text="Guida all'uso", command=self.guida_uso
        )
        self.info_btn.pack(pady=50, fill="x")

        # Oggetto
        self.subject_label = ctk.CTkLabel(self, text="Oggetto Email")
        self.subject_label.grid(row=0, column=1, sticky="w", padx=10, pady=(10, 0))
        self.subject_entry = ctk.CTkEntry(self)
        self.subject_entry.grid(row=1, column=1, sticky="ew", padx=10)
        # Corpo
        self.body_label = ctk.CTkLabel(
            self,
            text="Corpo Email (usa {{NomeColonna}} per dichiarare le colonne Excel)",
        )
        self.body_label.grid(row=2, column=1, sticky="nw", padx=10, pady=(10, 0))

        self.body_text = ctk.CTkTextbox(self)
        self.body_text.grid(row=3, column=1, sticky="nsew", padx=10, pady=(0, 10))

        # Progress
        self.progress = ctk.CTkProgressBar(self)
        self.progress.grid(row=4, column=1, sticky="ew", padx=10, pady=5)
        self.progress.set(0)

        self.status_label = ctk.CTkLabel(self, text="Nessun file caricato", anchor="w")
        self.status_label.grid(row=5, column=1, sticky="ew", padx=10, pady=(0, 10))

    # --- Funzioni ---
    def show_dialog(
        self, title: str, message: str, width: int = 500, height: int = 250
    ) -> None:
        dialog = ctk.CTkToplevel(self)
        dialog.title(title)
        dialog.geometry(f"{width}x{height}")
        dialog.grab_set()

        dialog.grid_columnconfigure(0, weight=1)
        dialog.grid_rowconfigure(0, weight=1)

        textbox = ctk.CTkTextbox(dialog, wrap="word")
        textbox.grid(row=0, column=0, sticky="nsew", padx=20, pady=20)
        textbox.insert("1.0", message)
        textbox.configure(state="disabled")

        close_btn = ctk.CTkButton(dialog, text="Chiudi", command=dialog.destroy)
        close_btn.grid(row=1, column=0, pady=(0, 15))

    def load_file(self) -> None:
        path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if not path:
            return
        try:
            self.df = pd.read_excel(path)
            self.file_path = path
            self.status_label.configure(text=f"File caricato: {os.path.basename(path)}")
        except Exception as e:
            self.show_dialog("Errore", f"Errore caricamento Excel:\n{e}")

    def load_word_template(self) -> None:
        path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
        if not path:
            return
        try:
            doc = Document(path)
            # keep all paragraphs (including empty) to preserve raw HTML/text
            full_text = [para.text for para in doc.paragraphs]
            if all((p is None or p == "") for p in full_text):
                self.show_dialog("Errore", "Il file Word è vuoto.")
                return
            subject = full_text[0] if len(full_text) > 0 else ""
            # join with single newline to preserve HTML structure as entered
            body = "\n".join(full_text[1:]) if len(full_text) > 1 else ""

            self.subject_entry.delete(0, "end")
            self.subject_entry.insert(0, subject)

            self.body_text.delete("1.0", "end")
            self.body_text.insert("1.0", body)

            self.show_dialog(
                "Template caricato", "Oggetto e corpo caricati correttamente."
            )
        except Exception as e:
            self.show_dialog("Errore", f"Errore caricamento Word:\n{e}")

    def create_word_template(self) -> None:
        try:
            doc = Document()
            doc.add_paragraph("Promemoria pagamento {{AnnoCorso}}")  # Oggetto
            doc.add_paragraph("")  # Riga vuota
            doc.add_paragraph(
                """<p>Gentili Genitori,</p>

            <p>Con la presente desideriamo ricordarvi il pagamento della retta relativa al secondo trimestre del corso di musica {{AnnoCorso}} frequentato da vostro/a figlio/a.</p>

            <p><b>L’importo dovuto è pari a {{Prezzo}} €</b> e può essere versato tramite bonifico bancario:</p>
            <ul>
            <li><b>IBAN: IT25L0863165011066000001528</b></li>
            <li><b>Beneficiario: Associazione Filarmonica Sanvitese</b></li>
            </ul>

            <p>Cogliamo inoltre l’occasione per ricordare che è necessario rinnovare la quota associativa e assicurativa, per un importo complessivo di <u>{{QuotaAssociativa}}</u> €.</p>

            <p>Tali quote dovranno essere versate presso la segreteria nei giorni di mercoledì e venerdì, dalle ore 16.30 alle ore 18.30, con pagamento in contanti entro il mese di {{MesePagamento}} {{AnnoPagamento}}.</p>

            <p>Restiamo a disposizione per eventuali chiarimenti e ringraziamo per la collaborazione.</p>

            <hr>

            <p><b>Cordiali saluti,</b><br>
            <b>Segreteria AFIS</b><br>
            <i>E. Pitton</i></p>

            <p>_____________________________</p>

            <p><b>Filarmonica Sanvitese APS</b><br>
            Piazzale Hermann Zotti, 1<br>
            33078 San Vito Al Tagliamento (PN)<br>
            Tel.: 3396771611<br>
            <a href="https://www.afisanvitese.it">www.afisanvitese.it</a><br>
            <a href="mailto:scuola@afisanvitese.it">scuola@afisanvitese.it</a></p>"""
            )

            save_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                initialfile="Matrice_mail.docx",
            )
            if save_path:
                doc.save(save_path)
                self.show_dialog("Creato", f"Template Word salvato come:\n{save_path}")
        except Exception as e:
            self.show_dialog("Errore", f"Errore creazione Word:\n{e}")

    def create_excel_template(self) -> None:
        try:
            # Crea workbook
            wb = openpyxl.Workbook()
            ws: Any = wb.active
            ws.title = "Template"

            # Intestazioni
            ws.append(["Email", "Nome", "Cognome", "AltroCampo"])

            # Riga esempio
            ws.append(["esempio@email.com", "Mario", "Rossi", "Valore"])

            save_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                initialfile="Matrice_mail.xlsx",
            )
            if save_path:
                wb.save(save_path)
                self.show_dialog("Creato", f"Template Excel salvato come:\n{save_path}")
        except Exception as e:
            self.show_dialog("Errore", f"Errore creazione Excel:\n{e}")

    def show_fields(self) -> None:
        if self.df is None:
            self.show_dialog("Attenzione", "Caricare prima un file Excel.")
            return
        columns = self.df.columns.tolist()
        # exclude the email column from template placeholders if it's named 'Email'
        placeholders = [col for col in columns if col.lower() != "email"]
        fields = "\n".join([f"{{{{{col}}}}}" for col in placeholders])
        self.show_dialog("Campi disponibili", fields)

    def guida_uso(self) -> None:

        guida: str = """
        
        1) Caricare un file excel dove ogni colonna corrisponde ad un campo
        
        2) Caricare o scrivere un testo in html (usare tranquillamente chat gpt per farlo generare)
        
        3) Per inserire (se non già fatto da chat gpt) i vari campi nell'oggetto e nel corpo che variano ad ogni singola mail inviata utilizzare {{NomeColonna}}
        
        4) Una volta completata la mail (oggetto e corpo) verranno inviate ad ogni mail presente nella prima colonna del file Excel
        
        5) Usare la modalità Test per inviare una mail di prova a se stessi
        
        6) Usare anteprima per generare una anteprima della prima mail
        
        7) Controllare il file email_log.txt per verficare il motivo di errori nell'invio di una mail
        
        """
        self.show_dialog("Campi disponibili", message=guida, width=1000, height=400)

    def validate_placeholders(self, text: str) -> Tuple[bool, Optional[List[str]]]:
        assert self.df is not None
        placeholders = re.findall(r"\{\{(.*?)\}\}", text)
        columns: List[str] = [c for c in self.df.columns.tolist()]
        missing: List[str] = []
        for ph in placeholders:
            key = ph.strip()
            if key not in columns:
                missing.append(key)
        if missing:
            return False, missing
        return True, None

    def preview_email(self) -> None:
        """Genera un'anteprima HTML dell'oggetto e del corpo usando la prima riga del DataFrame

        Se il corpo contiene tag HTML verrà trattato come HTML, altrimenti viene escapatto e
        i newline convertiti in <br>.
        """
        if self.df is None:
            self.show_dialog("Errore", "Caricare prima un file Excel.")
            return

        subject_template = self.subject_entry.get()
        body_template = self.body_text.get("1.0", "end")
        row = self.df.iloc[0]

        # Sostituisci placeholder in modo sicuro (NaN -> '')
        subject = subject_template
        body = body_template
        for col in self.df.columns:
            val = row[col]
            val_str = "" if pd.isna(val) else str(val)
            subject = subject.replace(f"{{{{{col}}}}}", val_str)
            body = body.replace(f"{{{{{col}}}}}", val_str)

        # Escape subject, body: if body seems to contain HTML keep it as-is, altrimenti escape + nl2br
        safe_subject = _html.escape(subject)
        body_is_html = "<" in body and ">" in body
        if body_is_html:
            safe_body = body
        else:
            safe_body = _html.escape(body).replace("\n", "<br>\n")

        html_content = f"""<!doctype html>
<html>
  <head>
    <meta charset=\"utf-8\"> 
    <title>Anteprima Email</title>
    <style>body{{font-family: Arial, Helvetica, sans-serif; padding:20px}} h2{{color:#333}}</style>
  </head>
  <body>
    <h2>{safe_subject}</h2>
    <hr>
    <div>{safe_body}</div>
  </body>
</html>
"""

        # scrivi su file temporaneo e apri nel browser predefinito
        try:
            tf = tempfile.NamedTemporaryFile(
                delete=False, suffix=".html", mode="w", encoding="utf-8"
            )
            tf.write(html_content)
            tf.flush()
            tf.close()
            webbrowser.open_new_tab(Path(tf.name).as_uri())
        except Exception as e:
            self.show_dialog(
                "Errore Anteprima", f"Impossibile creare anteprima HTML:\n{e}"
            )

    def send_emails(self) -> None:
        if self.df is None:
            self.show_dialog("Errore", "Caricare prima un file Excel.")
            return

        subject_template = self.subject_entry.get().strip()
        body_template = self.body_text.get("1.0", "end").strip()
        if not subject_template or not body_template:
            self.show_dialog("Errore", "Oggetto e corpo obbligatori.")
            return

        valid, wrong_field = self.validate_placeholders(
            subject_template + body_template
        )
        if not valid:
            missing = wrong_field or []
            self.show_dialog(
                "Errore Placeholder", f"Campi non trovati: {', '.join(missing)}"
            )
            return

        # disable send button and run sending in background thread to avoid UI freeze
        self.send_btn.configure(state="disabled")
        threading.Thread(
            target=self._send_emails_worker,
            args=(subject_template, body_template),
            daemon=True,
        ).start()

    def _is_valid_email(self, addr: str) -> bool:
        name, email = parseaddr(addr)
        if not email:
            return False
        return (
            re.match(r"^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$", email)
            is not None
        )

    def _validate_smtp_config(self) -> Tuple[bool, Optional[str]]:
        if not SMTP_SERVER:
            return False, "SMTP_SERVER non impostato"
        if not SMTP_USER:
            return False, "SMTP_USER non impostato"
        if not SMTP_PASSWORD:
            return False, "SMTP_PASSWORD non impostato"
        try:
            int(SMTP_PORT)
        except Exception:
            return False, "SMTP_PORT non valido"
        return True, None

    def _send_emails_worker(self, subject_template: str, body_template: str) -> None:
        ok, msg = self._validate_smtp_config()
        if not ok:
            self.after(0, lambda m=msg: self.show_dialog("Errore SMTP", m or ""))
            self.after(0, lambda: self.send_btn.configure(state="normal"))
            return

        server = None
        sent = 0
        errors = 0

        try:
            # mypy/static check: ensure these globals are not None
            assert (
                SMTP_SERVER is not None
                and SMTP_USER is not None
                and SMTP_PASSWORD is not None
            )
            server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
            server.ehlo()
            if USE_TLS:
                server.starttls()
                server.ehlo()
            server.login(SMTP_USER, SMTP_PASSWORD)
        except Exception as e:
            logging.error(f"Errore di connessione SMTP: {e}")
            self.after(
                0,
                lambda: self.show_dialog("Errore SMTP", f"Errore di connessione:\n{e}"),
            )
            self.after(0, lambda: self.send_btn.configure(state="normal"))
            if server:
                try:
                    server.quit()
                except Exception:
                    pass
            return

        assert self.df is not None
        total = 1 if self.test_mode.get() else len(self.df)
        index_counter = 0

        for i, (_, row) in enumerate(self.df.iterrows(), start=1):
            if self.test_mode.get() and i > 1:
                break

            index_counter += 1
            recipient = str(row.iloc[0]).strip() if len(row) > 0 else ""
            if self.test_mode.get():
                recipient = SMTP_USER

            if not self._is_valid_email(recipient):
                errors += 1
                logging.info(f"Email non valida: {recipient}")
                self.after(
                    0,
                    lambda r=recipient: self.status_label.configure(
                        text=f"Email non valida: {r}"
                    ),
                )
                self.after(0, lambda: self.progress.set(index_counter / total))
                continue

            # prepare subject and body replacing NaN safely
            subject = subject_template
            body = body_template
            for col in self.df.columns:
                val = row[col]
                if pd.isna(val):
                    val_str = ""
                else:
                    val_str = str(val)
                subject = subject.replace(f"{{{{{col}}}}}", val_str)
                body = body.replace(f"{{{{{col}}}}}", val_str)

            msg = MIMEMultipart()
            msg["From"] = SMTP_USER
            msg["To"] = recipient
            msg["Subject"] = subject
            msg.attach(MIMEText(body, "html"))

            try:
                server.send_message(msg)
                sent += 1
                logging.info(f"Email inviata correttamente a {recipient}")
            except Exception as e:
                errors += 1
                logging.error(f"Errore invio a {recipient}: {e}")

            self.after(0, lambda i=index_counter, t=total: self.progress.set(i / t))

        try:
            if server:
                server.quit()
        except Exception:
            pass

        self.after(0, lambda: self.send_btn.configure(state="normal"))
        self.after(
            0,
            lambda: self.show_dialog(
                "Invio completato",
                f"Email inviate: {sent}\nErrori: {errors}",
                width=550,
                height=300,
            ),
        )

    def _remove_preview_file(self, path: Path) -> None:
        try:
            if path and path.exists():
                path.unlink()
        except Exception:
            pass
        finally:
            if self._last_preview_path == path:
                self._last_preview_path = None

    def _schedule_preview_cleanup(self, path: Path, delay: int = 300) -> None:
        try:
            t = threading.Timer(delay, lambda: self._remove_preview_file(path))
            t.daemon = True
            t.start()
        except Exception:
            pass


if __name__ == "__main__":
    app = EmailApp()
    app.mainloop()
